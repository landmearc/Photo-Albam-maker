import os
import re
import copy
import math
from dataclasses import dataclass
from datetime import datetime
from typing import List, Tuple, Optional, Dict

import pdfplumber
from PIL import Image, ImageOps, ExifTags, ImageTk, ImageDraw

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from tkinterdnd2 import DND_FILES, TkinterDnD

from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

import fitz  # PyMuPDF


# -------------------------
# Settings / Const
# -------------------------
IMG_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}
TARGET_W = 640
TARGET_H = 480

# 半角英数系
RE_BOUNDARY = re.compile(r"^[A-Z]{1,4}\d{1,4}$")         # PK1, BYO12, PL3, CO1
RE_TPOINT = re.compile(r"^T\d{1,4}$")                    # T1, T01
RE_A1 = re.compile(r"^[A-Z]\d{1,4}$")                    # P1, K12
RE_CTRL_PAIR_A = re.compile(r"^[A-Z0-9]{2,12}$")
RE_CTRL_PAIR_B = re.compile(r"^\d{1,4}-\d{1,4}$")

# 数値・地番系
RE_DECIMAL = re.compile(r"^-?\d+\.\d+$")
RE_INTLIKE = re.compile(r"^-?\d+$")
RE_LOT_HYPHEN = re.compile(r"^\d{3,6}-\d{1,4}$")         # 8844-1
RE_LOT_INT = re.compile(r"^\d{3,6}$")                    # 8844
RE_NUM_HYPHEN_NUM = re.compile(r"^\d{1,4}-\d{1,4}$")     # 12-3
RE_COORDLIKE = re.compile(r"^-?\d+(\.\d+)?$")

# 日本語点名系
RE_JP_POINT_1 = re.compile(r"^[ぁ-んァ-ヶー一-龠]+(?:杭)?\d{1,4}$")                 # ブラ杭1, コン杭5
RE_JP_POINT_2 = re.compile(r"^[ぁ-んァ-ヶー一-龠]+[A-Z]{1,4}\d{1,4}$")             # 市PK2
RE_JP_POINT_3 = re.compile(r"^[A-Z]{1,4}[ぁ-んァ-ヶー一-龠]+\d{1,4}$")             # PK杭2
RE_JP_POINT_4 = re.compile(r"^[ぁ-んァ-ヶー一-龠A-Z0-9]+-\d{1,4}$")                # 稀な派生
RE_HEADER_WORDS = {
    "求積表", "座標一覧", "点名", "番号", "測点", "辺長", "面積", "地番",
    "北距", "東距", "X", "Y", "XN", "YN", "X座標", "Y座標", "備考", "NO"
}

PDF_POINT_PATTERNS = [
    RE_BOUNDARY,
    RE_TPOINT,
    RE_A1,
    RE_JP_POINT_1,
    RE_JP_POINT_2,
    RE_JP_POINT_3,
    RE_JP_POINT_4,
]

FULLWIDTH_TRANS = str.maketrans({
    "０": "0", "１": "1", "２": "2", "３": "3", "４": "4",
    "５": "5", "６": "6", "７": "7", "８": "8", "９": "9",

    "Ａ": "A", "Ｂ": "B", "Ｃ": "C", "Ｄ": "D", "Ｅ": "E",
    "Ｆ": "F", "Ｇ": "G", "Ｈ": "H", "Ｉ": "I", "Ｊ": "J",
    "Ｋ": "K", "Ｌ": "L", "Ｍ": "M", "Ｎ": "N", "Ｏ": "O",
    "Ｐ": "P", "Ｑ": "Q", "Ｒ": "R", "Ｓ": "S", "Ｔ": "T",
    "Ｕ": "U", "Ｖ": "V", "Ｗ": "W", "Ｘ": "X", "Ｙ": "Y",
    "Ｚ": "Z",

    "ａ": "a", "ｂ": "b", "ｃ": "c", "ｄ": "d", "ｅ": "e",
    "ｆ": "f", "ｇ": "g", "ｈ": "h", "ｉ": "i", "ｊ": "j",
    "ｋ": "k", "ｌ": "l", "ｍ": "m", "ｎ": "n", "ｏ": "o",
    "ｐ": "p", "ｑ": "q", "ｒ": "r", "ｓ": "s", "ｔ": "t",
    "ｕ": "u", "ｖ": "v", "ｗ": "w", "ｘ": "x", "ｙ": "y",
    "ｚ": "z",

    "（": "(", "）": ")",
    "［": "[", "］": "]",
    "｛": "{", "｝": "}",
    "＋": "+",
    "－": "-",
    "ー": "-",
    "―": "-",
    "‐": "-",
    "−": "-",
    "．": ".",
    "，": ",",
    "　": " ",
})

# -------------------------
# Utility
# -------------------------
def safe_mkdir(p: str) -> None:
    if p:
        os.makedirs(p, exist_ok=True)


def is_image_file(p: str) -> bool:
    return os.path.splitext(p.lower())[1] in IMG_EXTS


def norm_paths_from_dnd(data: str) -> List[str]:
    if not data:
        return []
    paths = []
    buff = ""
    in_brace = False
    for ch in data:
        if ch == "{":
            in_brace = True
            buff = ""
        elif ch == "}":
            in_brace = False
            if buff:
                paths.append(buff)
                buff = ""
        elif ch == " " and not in_brace:
            if buff:
                paths.append(buff)
                buff = ""
        else:
            buff += ch
    if buff:
        paths.append(buff)
    return [p.strip().strip('"') for p in paths]


def resize_keep_aspect(src_path: str, dst_path: str, target_w: int = TARGET_W, target_h: int = TARGET_H) -> Tuple[int, int]:
    with Image.open(src_path) as im:
        im = ImageOps.exif_transpose(im)
        w, h = im.size
        if h == 0:
            raise ValueError("Invalid image height")

        ratio = w / h
        is_4_3ish = abs(ratio - (4 / 3)) < 0.03

        if is_4_3ish:
            scale = min(target_w / w, target_h / h)
        else:
            scale = target_h / h

        new_w = max(1, int(round(w * scale)))
        new_h = max(1, int(round(h * scale)))

        im2 = im.resize((new_w, new_h), Image.LANCZOS)

        safe_mkdir(os.path.dirname(dst_path))
        ext = os.path.splitext(dst_path.lower())[1]
        if ext in (".jpg", ".jpeg"):
            im2.save(dst_path, quality=90, optimize=True)
        else:
            im2.save(dst_path)
        return new_w, new_h


def get_exif_datetime_key(path: str):
    dt = None
    sub = 0

    try:
        with Image.open(path) as im:
            exif = im.getexif()
            if exif:
                tag_map = {ExifTags.TAGS.get(k, k): v for k, v in exif.items()}
                for key in ("DateTimeOriginal", "DateTimeDigitized", "DateTime"):
                    if key in tag_map:
                        s = str(tag_map[key])
                        try:
                            dt = datetime.strptime(s, "%Y:%m:%d %H:%M:%S")
                        except Exception:
                            dt = None
                        break
                if dt is not None:
                    ssub = str(tag_map.get("SubSecTimeOriginal", "0"))
                    ssub = re.sub(r"\D", "", ssub) or "0"
                    try:
                        sub = int(ssub)
                    except Exception:
                        sub = 0
    except Exception:
        pass

    if dt is None:
        try:
            dt = datetime.fromtimestamp(os.path.getmtime(path))
            sub = 0
        except Exception:
            dt = datetime.min
            sub = 0

    return (dt, sub, os.path.basename(path).lower())


def _to_wareki(dt: datetime) -> str:
    y, m, d = dt.year, dt.month, dt.day
    if (y, m, d) >= (2019, 5, 1):
        return f"令和{y - 2018}年{m}月{d}日"
    elif (y, m, d) >= (1989, 1, 8):
        return f"平成{y - 1988}年{m}月{d}日"
    elif (y, m, d) >= (1926, 12, 25):
        return f"昭和{y - 1925}年{m}月{d}日"
    else:
        return f"{y}年{m}月{d}日"


def get_shoot_date_jp(path: str) -> str:
    dt = None
    try:
        with Image.open(path) as im:
            exif = im.getexif()
            if exif:
                tag_map = {ExifTags.TAGS.get(k, k): v for k, v in exif.items()}
                for key in ("DateTimeOriginal", "DateTimeDigitized", "DateTime"):
                    if key in tag_map:
                        s = str(tag_map[key])
                        dt = datetime.strptime(s, "%Y:%m:%d %H:%M:%S")
                        break
    except Exception:
        dt = None

    if dt is None:
        try:
            dt = datetime.fromtimestamp(os.path.getmtime(path))
        except Exception:
            return ""

    return f"撮影日：{_to_wareki(dt)}"


def normalize_point(s: str) -> str:
    s = (s or "").strip()
    s = s.translate(FULLWIDTH_TRANS)
    s = s.replace("　", " ")
    s = re.sub(r"\s+", "", s)
    s = s.upper()
    return s


def to_fullwidth_digits(s: str) -> str:
    return s.translate(str.maketrans("0123456789", "０１２３４５６７８９"))


def _is_noise_token(tok: str) -> bool:
    t = normalize_point(tok)
    if not t:
        return True
    if t in {w.upper() for w in RE_HEADER_WORDS}:
        return True
    if RE_DECIMAL.fullmatch(t):
        return True
    if RE_INTLIKE.fullmatch(t):
        return True
    if RE_LOT_HYPHEN.fullmatch(t):
        return True
    if RE_LOT_INT.fullmatch(t):
        return True
    if RE_NUM_HYPHEN_NUM.fullmatch(t):
        return True
    return False


def _dedupe_keep_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in items:
        if x in seen:
            continue
        seen.add(x)
        out.append(x)
    return out


def _cell_text(x) -> str:
    if x is None:
        return ""
    return str(x).replace("\n", " ").strip()


def is_point_token(tok: str) -> bool:
    t = normalize_point(tok)
    if not t:
        return False
    if _is_noise_token(t):
        return False
    for pat in PDF_POINT_PATTERNS:
        if pat.fullmatch(t):
            return True
    return False


def _clean_pdf_token(tok: str) -> str:
    t = (tok or "").strip()
    t = t.translate(FULLWIDTH_TRANS)
    t = t.replace("　", " ")
    t = t.replace("．", ".")
    t = t.replace("，", ",")
    t = t.replace("―", "-")
    t = t.replace("‐", "-")
    t = t.replace("-", "-")
    t = t.replace("−", "-")
    t = re.sub(r"^[\[\(\{<＜〈《「『【]+", "", t)
    t = re.sub(r"[\]\)\}>＞〉》」』】,.;:]+$", "", t)
    return normalize_point(t)


def _table_has_coord_header(table: List[List[str]]) -> bool:
    for row in table[:6]:
        r = " ".join(_cell_text(c) for c in row)
        ru = normalize_point(r)
        if ("点名" in ru or "測点" in ru) and (("X" in ru) or ("XN" in ru)) and (("Y" in ru) or ("YN" in ru)):
            return True
    return False


def _table_has_area_header(table: List[List[str]]) -> bool:
    for row in table[:8]:
        r = " ".join(_cell_text(c) for c in row)
        ru = normalize_point(r)
        if "求積表" in ru:
            return True
        if ("NO" in ru) and (("XN" in ru) or ("YN" in ru) or (("X" in ru) and ("Y" in ru))):
            return True
        if ("点名" in ru or "測点" in ru) and ("辺長" in ru or "面積" in ru):
            return True
    return False


def _extract_points_from_table_anywhere(table: List[List[str]]) -> List[str]:
    out = []
    for row in table:
        for cell in row:
            txt = _cell_text(cell)
            if not txt:
                continue
            parts = re.split(r"[\s,]+", txt)
            for p in parts:
                q = _clean_pdf_token(p)
                if is_point_token(q):
                    out.append(q)
    return out


def _extract_points_from_text_block(text: str) -> List[str]:
    out = []
    if not text:
        return out

    lines = text.splitlines()
    in_area = False
    in_coord = False

    for line in lines:
        s = (line or "").strip()
        if not s:
            continue

        su = normalize_point(s)

        if "求積表" in su:
            in_area = True
            in_coord = False
            continue

        if ("点名" in su or "測点" in su) and (("X" in su) or ("XN" in su)) and (("Y" in su) or ("YN" in su)):
            in_coord = True
            in_area = False
            continue

        tokens = re.split(r"[\s,]+", s)
        cand = [_clean_pdf_token(t) for t in tokens if t.strip()]

        if in_coord or in_area:
            for t in cand:
                if is_point_token(t):
                    out.append(t)
        else:
            for t in cand:
                if is_point_token(t):
                    out.append(t)

    return out


def _extract_points_from_pdfplumber_words(page) -> List[str]:
    out = []
    try:
        words = page.extract_words(
            keep_blank_chars=False,
            use_text_flow=True,
            extra_attrs=[]
        )
    except Exception:
        words = []

    for w in words:
        txt = _clean_pdf_token(w.get("text", ""))
        if is_point_token(txt):
            out.append(txt)
    return out


def _extract_points_from_fitz_page(fitz_page) -> List[str]:
    out = []

    try:
        words = fitz_page.get_text("words")
        for w in words:
            txt = _clean_pdf_token(w[4])
            if is_point_token(txt):
                out.append(txt)
    except Exception:
        pass

    try:
        txt = fitz_page.get_text("text")
        out.extend(_extract_points_from_text_block(txt))
    except Exception:
        pass

    return out


def extract_point_names_from_pdf(pdf_path: str) -> List[str]:
    """
    求積表・座標一覧の点名を漏れにくく抽出する
    - pdfplumber table
    - pdfplumber words
    - PyMuPDF words/text
    を併用
    """
    points_priority: List[str] = []
    points_general: List[str] = []

    plumber_pdf = None
    fitz_pdf = None

    try:
        plumber_pdf = pdfplumber.open(pdf_path)
        fitz_pdf = fitz.open(pdf_path)

        page_count = min(len(plumber_pdf.pages), len(fitz_pdf))

        for i in range(page_count):
            p_page = plumber_pdf.pages[i]
            f_page = fitz_pdf[i]

            try:
                tables = p_page.extract_tables(
                    table_settings={
                        "vertical_strategy": "lines",
                        "horizontal_strategy": "lines",
                        "intersection_tolerance": 5,
                        "snap_tolerance": 3,
                        "join_tolerance": 3,
                        "edge_min_length": 6,
                        "min_words_vertical": 1,
                        "min_words_horizontal": 1,
                    }
                )
            except Exception:
                tables = []

            for tbl in tables or []:
                if not tbl:
                    continue
                table = [[_cell_text(c) for c in row] for row in tbl]
                pts = _extract_points_from_table_anywhere(table)
                if _table_has_coord_header(table) or _table_has_area_header(table):
                    points_priority.extend(pts)
                else:
                    points_general.extend(pts)

            try:
                txt = p_page.extract_text() or ""
            except Exception:
                txt = ""
            points_priority.extend(_extract_points_from_text_block(txt))

            points_general.extend(_extract_points_from_pdfplumber_words(p_page))
            points_general.extend(_extract_points_from_fitz_page(f_page))

    finally:
        try:
            if plumber_pdf:
                plumber_pdf.close()
        except Exception:
            pass
        try:
            if fitz_pdf:
                fitz_pdf.close()
        except Exception:
            pass

    merged = _dedupe_keep_order(
        [normalize_point(x) for x in (points_priority + points_general) if normalize_point(x)]
    )

    final_out = []
    for t in merged:
        if is_point_token(t):
            final_out.append(t)

    return final_out


# -------------------------
# SIMA parser
# -------------------------
def read_text_guess_encoding(path: str) -> str:
    for enc in ("cp932", "shift_jis", "utf-8"):
        try:
            return open(path, "r", encoding=enc).read()
        except Exception:
            continue
    return open(path, "r", encoding="utf-8", errors="replace").read()


def parse_sima_file(path: str) -> Tuple[Dict[str, Tuple[float, float, Optional[float]]], List[List[str]]]:
    txt = read_text_guess_encoding(path)
    lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]

    points: Dict[str, Tuple[float, float, Optional[float]]] = {}
    polylines: List[List[str]] = []

    in_d_block = False
    current_poly: List[str] = []

    for ln in lines:
        parts = [p.strip() for p in ln.split(",")]
        if not parts:
            continue

        head = parts[0].upper()

        if head == "A01":
            if len(parts) >= 5:
                name = (parts[2] if len(parts) > 2 else "").strip()
                if name:
                    try:
                        x = float(parts[3])  # 北方向
                        y = float(parts[4])  # 東西方向
                        z = None
                        if len(parts) >= 6 and parts[5] != "":
                            try:
                                z = float(parts[5])
                            except Exception:
                                z = None
                        points[name] = (x, y, z)
                    except Exception:
                        pass
            continue

        if head == "D00":
            in_d_block = True
            current_poly = []
            continue

        if head == "D99":
            if in_d_block and len(current_poly) >= 2:
                polylines.append(current_poly[:])
            in_d_block = False
            current_poly = []
            continue

        if in_d_block and head == "B01":
            if len(parts) >= 3:
                name = parts[2].strip()
                if name:
                    current_poly.append(name)
            continue

    return points, polylines


# -------------------------
# Word album (direct layout)
# -------------------------
def _clear_cell_paragraphs(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    cell.add_paragraph("")


def _set_cell_border(cell, border_size_pt: int = 4):
    """セルの罫線を設定する"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size_pt * 8))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)



def _set_row_height_exact(row, height_mm: float):
    """行高さを exact で固定する"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    twips = int(height_mm * 56.7)
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _set_page_break_before(paragraph):
    """段落に改ページ前を設定する"""
    pPr = paragraph._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "1")
    pPr.append(pb)


def build_album_direct(
    output_docx: str,
    rows: List[Tuple[str, str, str, str, str]],
    photos_per_page: int = 4,
    img_width_mm: int = 75,
):
    """
    テンプレート不要で写真帳docxを直接生成。
    ページごとに独立テーブル。余分な空段落を除去して空白ページを防ぐ。
    rows: (img_path, name_no_ext, shoot_date_str, category, note)
    """
    from docx.oxml.ns import qn as _qn2

    def _remove_para(para):
        """段落要素をドキュメントから削除する"""
        p = para._element
        p.getparent().remove(p)

    doc = Document()

    # ドキュメント冒頭のデフォルト空段落を削除
    for p in list(doc.paragraphs):
        _remove_para(p)

    section = doc.sections[0]
    section.page_width    = Mm(210)
    section.page_height   = Mm(297)
    section.left_margin   = Mm(20)
    section.right_margin  = Mm(20)
    section.top_margin    = Mm(15)
    section.bottom_margin = Mm(15)

    TOTAL_W  = Mm(170)
    PHOTO_W  = Mm(img_width_mm)
    TEXT_W   = TOTAL_W - PHOTO_W
    ROW_H_MM = (267.0 - 2) / photos_per_page
    IMG_H_MM = ROW_H_MM - 5

    style = doc.styles["Normal"]
    style.font.name = "MS 明朝"
    style.font.size = Pt(9)

    pages = [rows[i:i + photos_per_page] for i in range(0, len(rows), photos_per_page)]

    for page_idx, page_rows in enumerate(pages):

        # ページ2以降：pageBreak run を持つ最小段落をテーブルの前に挿入
        if page_idx > 0:
            pb_para = OxmlElement("w:p")
            pb_pPr  = OxmlElement("w:pPr")
            pb_pSp  = OxmlElement("w:jc")   # dummy
            pb_r    = OxmlElement("w:r")
            pb_br   = OxmlElement("w:br")
            pb_br.set(_qn2("w:type"), "page")
            pb_r.append(pb_br)
            pb_para.append(pb_r)
            # body の末尾に追加
            doc.element.body.append(pb_para)

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        for img_path, name_no_ext, shoot_date_str, category, note in page_rows:

            row = table.add_row()

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(_qn2("w:val"), str(int(ROW_H_MM * 56.7)))
            trHeight.set(_qn2("w:hRule"), "exact")
            trPr.append(trHeight)

            row.cells[0].width = PHOTO_W
            row.cells[1].width = TEXT_W

            # 左セル：写真
            c0 = row.cells[0]
            _clear_cell_paragraphs(c0)
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after  = Pt(0)

            tcPr0 = c0._tc.get_or_add_tcPr()
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(_qn2("w:val"), "center")
            tcPr0.append(vAlign)

            run0 = p0.add_run()
            try:
                with Image.open(img_path) as im:
                    iw, ih = im.size
                ratio = iw / max(ih, 1)
                calc_w = IMG_H_MM * ratio
                if calc_w > img_width_mm - 2:
                    calc_w = img_width_mm - 2
                run0.add_picture(img_path, width=Mm(calc_w))
            except Exception:
                p0.add_run(f"[画像読込失敗] {os.path.basename(img_path)}")

            _set_cell_border(c0)

            # 右セル：テキスト
            c1 = row.cells[1]
            _clear_cell_paragraphs(c1)

            p_name = c1.paragraphs[0]
            p_name.paragraph_format.space_before = Pt(5)
            p_name.paragraph_format.space_after  = Pt(2)
            rn = p_name.add_run(name_no_ext)
            rn.bold = True
            rn.underline = True
            rn.font.size = Pt(10)

            p_date = c1.add_paragraph(shoot_date_str)
            p_date.paragraph_format.space_after = Pt(2)

            p_cat = c1.add_paragraph(category)
            p_cat.paragraph_format.space_after = Pt(2)

            p_note = c1.add_paragraph(note or "")
            p_note.paragraph_format.space_after = Pt(1)

            blank_count = max(2, int((ROW_H_MM - 40) / 7))
            for _ in range(blank_count):
                pb2 = c1.add_paragraph("")
                pb2.paragraph_format.space_before = Pt(0)
                pb2.paragraph_format.space_after  = Pt(0)

            _set_cell_border(c1)

        # テーブル直後に python-docx が自動追加する空段落を削除
        paras_after = doc.paragraphs
        if paras_after:
            _remove_para(paras_after[-1])

    safe_mkdir(os.path.dirname(output_docx) or ".")
    doc.save(output_docx)

# -------------------------
# Data model
# -------------------------
@dataclass
class PhotoItem:
    src_path: str
    resized_path: str = ""
    assigned_point: str = ""
    renamed_path: str = ""
    category: str = "新設"
    note: str = ""


# -------------------------
# App
# -------------------------
class App(TkinterDnD.Tk):
    def __init__(self):
        super().__init__()
        self.title("写真帳メーカー v25")
        self.geometry("1700x980")

        self.photos: List[PhotoItem] = []
        self.point_names: List[str] = []
        self.current_pdf: str = ""
        self.current_sim: str = ""

        self.category_var = tk.StringVar(value="新設")
        self.category_manual_var = tk.StringVar(value="新設")
        self.point_edit_var = tk.StringVar()
        self.preview_mode_var = tk.StringVar(value="pdf")
        self.output_name_var = tk.StringVar(value="写真帳")

        self.rename_mode_var = tk.StringVar(value="point")
        self.serial_prefix_var = tk.StringVar(value="")
        self.serial_start_var = tk.IntVar(value=1)
        self.serial_digits_var = tk.IntVar(value=2)
        self.serial_fullwidth_var = tk.BooleanVar(value=True)

        self._photo_preview_tk = None
        self._current_photo_preview_path = ""

        self._pdf_doc: Optional[fitz.Document] = None
        self._pdf_page_index = 0
        self._pdf_highlight_word = ""
        self._pdf_preview_tk = None
        self._pdf_zoom = 1.5
        self._pdf_min_zoom = 0.5
        self._pdf_max_zoom = 6.0
        self._pdf_pan_start = None
        self._pdf_offx = 20.0
        self._pdf_offy = 20.0
        self._pdf_cache = {}
        self._pdf_dragging = False
        self._pdf_canvas_initialized = False

        self.sima_points: Dict[str, Tuple[float, float, Optional[float]]] = {}
        self.sima_polylines: List[List[str]] = []
        self._sim_zoom = 1.0
        self._sim_min_zoom = 0.05
        self._sim_max_zoom = 50.0
        self._sim_offx = 0.0
        self._sim_offy = 0.0
        self._sim_bbox = None
        self._sim_selected = ""

        self.sim_pick_radius_var = tk.IntVar(value=10)
        self.sim_drag_threshold_var = tk.IntVar(value=6)
        self._sim_pan_start = None
        self._sim_pan_origin = None
        self._sim_pan_moved = False

        self._build_ui()

    def logln(self, s: str):
        print(s)

    def _event_has_ctrl(self, event) -> bool:
        state = getattr(event, "state", 0)
        return bool(state & 0x0004) or bool(state & 0x0008)

    def _event_has_shift(self, event) -> bool:
        state = getattr(event, "state", 0)
        return bool(state & 0x0001)

    def _build_ui(self):
        root = ttk.Frame(self, padding=8)
        root.pack(fill="both", expand=True)

        main = ttk.Panedwindow(root, orient="horizontal")
        main.pack(fill="both", expand=True)

        # =============================================
        # LEFT with scroll
        # =============================================
        left_container = ttk.Frame(main)
        main.add(left_container, weight=2)

        self.left_canvas = tk.Canvas(left_container, highlightthickness=0)
        left_scroll_y = ttk.Scrollbar(left_container, orient="vertical", command=self.left_canvas.yview)
        self.left_canvas.configure(yscrollcommand=left_scroll_y.set)

        left_scroll_y.pack(side="right", fill="y")
        self.left_canvas.pack(side="left", fill="both", expand=True)

        left_inner = ttk.Frame(self.left_canvas, padding=6)
        self.left_canvas_window = self.left_canvas.create_window((0, 0), window=left_inner, anchor="nw")

        def _on_left_inner_configure(event):
            self.left_canvas.configure(scrollregion=self.left_canvas.bbox("all"))

        def _on_left_canvas_configure(event):
            self.left_canvas.itemconfig(self.left_canvas_window, width=event.width)

        left_inner.bind("<Configure>", _on_left_inner_configure)
        self.left_canvas.bind("<Configure>", _on_left_canvas_configure)

        self.left_canvas.bind_all("<MouseWheel>", self._on_global_mousewheel, add="+")
        self.left_canvas.bind_all("<Button-4>", self._on_global_mousewheel_linux, add="+")
        self.left_canvas.bind_all("<Button-5>", self._on_global_mousewheel_linux, add="+")

        left = left_inner

        ttk.Label(left, text="① 写真（ここにドラッグ＆ドロップ / または追加）").pack(anchor="w")

        drop_ph = ttk.Label(left, text="ここに写真ファイルをドロップ", relief="ridge", padding=10)
        drop_ph.pack(fill="x", pady=4)
        drop_ph.drop_target_register(DND_FILES)
        drop_ph.dnd_bind("<<Drop>>", self.on_drop_photos)

        btns = ttk.Frame(left)
        btns.pack(fill="x", pady=4)

        ttk.Button(btns, text="写真追加", command=self.add_photos_dialog).pack(side="left")
        ttk.Button(btns, text="一覧クリア", command=self.clear_photos).pack(side="left", padx=4)
        ttk.Button(btns, text="② EXIFで並び替え", command=self.sort_photos_by_exif).pack(side="left", padx=12)
        ttk.Button(btns, text="リサイズ（resize/へ）", command=self.do_resize).pack(side="left", padx=4)
        ttk.Button(btns, text="リネーム実行", command=self.do_rename_global_prefix).pack(side="left", padx=4)

        rename_opt = ttk.LabelFrame(left, text="リネーム設定", padding=6)
        rename_opt.pack(fill="x", pady=(0, 8))

        rowr1 = ttk.Frame(rename_opt)
        rowr1.pack(fill="x")
        ttk.Radiobutton(rowr1, text="点名付き", variable=self.rename_mode_var, value="point").pack(side="left")
        ttk.Radiobutton(rowr1, text="連番のみ", variable=self.rename_mode_var, value="serial").pack(side="left", padx=10)

        rowr2 = ttk.Frame(rename_opt)
        rowr2.pack(fill="x", pady=(4, 0))
        ttk.Label(rowr2, text="接頭語").pack(side="left")
        ttk.Entry(rowr2, textvariable=self.serial_prefix_var, width=8).pack(side="left", padx=4)
        ttk.Label(rowr2, text="開始番号").pack(side="left", padx=(12, 0))
        ttk.Spinbox(rowr2, from_=0, to=99999, textvariable=self.serial_start_var, width=6).pack(side="left", padx=4)
        ttk.Label(rowr2, text="桁数").pack(side="left", padx=(12, 0))
        ttk.Spinbox(rowr2, from_=1, to=10, textvariable=self.serial_digits_var, width=4).pack(side="left", padx=4)
        ttk.Checkbutton(rowr2, text="全角数字", variable=self.serial_fullwidth_var).pack(side="left", padx=12)

        photo_list_box = ttk.LabelFrame(left, text="写真一覧", padding=4)
        photo_list_box.pack(fill="x", pady=(0, 8))

        photo_list_box = ttk.Frame(left, relief="groove", borderwidth=2)
        photo_list_box.pack(fill="x", pady=(2, 8))

        tree_wrap = ttk.Frame(photo_list_box)
        tree_wrap.pack(fill="x", expand=False, padx=2, pady=2)

        cols = ("No", "元ファイル", "点名", "区分", "説明", "resize出力", "リネーム後")
        self.tree = ttk.Treeview(tree_wrap, columns=cols, show="headings", height=5)
        for c in cols:
            self.tree.heading(c, text=c)
        self.tree.column("No", width=40, anchor="e")
        self.tree.column("元ファイル", width=160)
        self.tree.column("点名", width=80)
        self.tree.column("区分", width=70)
        self.tree.column("説明", width=160)
        self.tree.column("resize出力", width=120)
        self.tree.column("リネーム後", width=120)

        tree_y = ttk.Scrollbar(tree_wrap, orient="vertical",   command=self.tree.yview)
        tree_x = ttk.Scrollbar(tree_wrap, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=tree_y.set, xscrollcommand=tree_x.set)

        tree_y.pack(side="right", fill="y")
        tree_x.pack(side="bottom", fill="x")
        self.tree.pack(side="left", fill="both", expand=True)
        self.tree.bind("<<TreeviewSelect>>", self.on_photo_select_preview)
        self.tree.bind("<Button-1>", self._on_tree_click, add="+")
        self.tree.bind("<Double-Button-1>", self._on_tree_dblclick, add="+")

        ttk.Label(left, text="③-1 測量図PDF（ドロップ）→ 点名抽出（求積表・座標一覧強化）").pack(anchor="w")
        drop_pdf = ttk.Label(left, text="ここにPDFをドロップ", relief="ridge", padding=10)
        drop_pdf.pack(fill="x", pady=4)
        drop_pdf.drop_target_register(DND_FILES)
        drop_pdf.dnd_bind("<<Drop>>", self.on_drop_pdf)

        self.pdf_label = ttk.Label(left, text="PDF: 未選択")
        self.pdf_label.pack(anchor="w", pady=2)

        rowpdf = ttk.Frame(left)
        rowpdf.pack(fill="x", pady=(0, 6))
        ttk.Button(rowpdf, text="PDF選択…", command=self.pick_pdf_dialog).pack(side="left")
        ttk.Button(rowpdf, text="PDFから点名抽出", command=self.do_extract_points_pdf).pack(side="left", padx=6)

        ttk.Label(left, text="③-2 SIMA(.sim/.sima)（ドロップ）→ 点名抽出＋表示").pack(anchor="w")
        drop_sim = ttk.Label(left, text="ここにSIMファイルをドロップ", relief="ridge", padding=10)
        drop_sim.pack(fill="x", pady=4)
        drop_sim.drop_target_register(DND_FILES)
        drop_sim.dnd_bind("<<Drop>>", self.on_drop_sim)

        self.sim_label = ttk.Label(left, text="SIMA: 未選択")
        self.sim_label.pack(anchor="w", pady=2)

        rowsim = ttk.Frame(left)
        rowsim.pack(fill="x", pady=(0, 6))
        ttk.Button(rowsim, text="SIM選択…", command=self.pick_sim_dialog).pack(side="left")
        ttk.Button(rowsim, text="SIMから点名抽出", command=self.do_extract_points_sim).pack(side="left", padx=6)
        ttk.Button(rowsim, text="SIM表示をフィット", command=self.sim_fit_view).pack(side="left", padx=6)

        ttk.Label(left, text="点名一覧（抽出結果）").pack(anchor="w")
        pts_frame = ttk.Frame(left)
        pts_frame.pack(fill="both", expand=True, pady=(2, 8))
        pts_y = ttk.Scrollbar(pts_frame, orient="vertical")
        pts_x = ttk.Scrollbar(pts_frame, orient="horizontal")
        self.points_list = tk.Listbox(pts_frame, height=10, selectmode="extended",
                                      yscrollcommand=pts_y.set, xscrollcommand=pts_x.set)
        pts_y.config(command=self.points_list.yview)
        pts_x.config(command=self.points_list.xview)
        pts_y.pack(side="right",  fill="y")
        pts_x.pack(side="bottom", fill="x")
        self.points_list.pack(side="left", fill="both", expand=True)
        self.points_list.bind("<<ListboxSelect>>", self.on_point_select)
        self.points_list.bind("<Double-Button-1>", self.on_point_double_click_assign)

        edit = ttk.LabelFrame(left, text="点名の手動修正", padding=6)
        edit.pack(fill="x", pady=(0, 8))

        rowe = ttk.Frame(edit)
        rowe.pack(fill="x")
        ttk.Label(rowe, text="点名:").pack(side="left")
        ttk.Entry(rowe, textvariable=self.point_edit_var, width=18).pack(side="left", padx=6)
        ttk.Button(rowe, text="追加", command=self.point_add).pack(side="left")
        ttk.Button(rowe, text="更新", command=self.point_update).pack(side="left", padx=4)
        ttk.Button(rowe, text="削除", command=self.point_delete).pack(side="left", padx=4)
        ttk.Button(rowe, text="↑", width=3, command=self.point_move_up).pack(side="right")
        ttk.Button(rowe, text="↓", width=3, command=self.point_move_down).pack(side="right", padx=4)

        assign = ttk.LabelFrame(left, text="④ 点名割当 / 並び替え", padding=6)
        assign.pack(fill="x", pady=(0, 8))
        ttk.Label(assign, text="点名ダブルクリック → 左の選択写真へ一括割当").pack(anchor="w")

        rowa = ttk.Frame(assign)
        rowa.pack(fill="x", pady=(4, 0))
        ttk.Button(rowa, text="選択分を順番に割当", command=self.assign_points_in_order).pack(side="left")
        ttk.Button(rowa, text="選択写真の点名クリア", command=self.clear_selected_assignments).pack(side="left", padx=6)


        catf = ttk.LabelFrame(left, text="⑤ 写真ごとの区分", padding=6)
        catf.pack(fill="x", pady=(0, 8))
        rowc = ttk.Frame(catf)
        rowc.pack(fill="x")
        ttk.Label(rowc, text="区分:").pack(side="left")
        cat_combo = ttk.Combobox(rowc, textvariable=self.category_var, values=["新設", "復元", "既設"], width=8, state="readonly")
        cat_combo.pack(side="left", padx=4)
        ttk.Button(rowc, text="選択値を入力欄へ", command=self.copy_category_choice_to_manual).pack(side="left", padx=4)
        ttk.Label(rowc, text="手入力:").pack(side="left", padx=(12, 4))
        ttk.Entry(rowc, textvariable=self.category_manual_var, width=12).pack(side="left")
        ttk.Button(rowc, text="選択写真へ区分設定", command=self.apply_category_to_selected).pack(side="left", padx=6)
        ttk.Button(rowc, text="全写真へ区分設定", command=self.apply_category_to_all).pack(side="left", padx=4)

        alb = ttk.LabelFrame(left, text="⑥ 写真帳（保存先＝写真フォルダ）", padding=6)
        alb.pack(fill="x", pady=(0, 8))
        row_alb = ttk.Frame(alb)
        row_alb.pack(fill="x", pady=(0, 4))
        ttk.Label(row_alb, text="DOCX名:").pack(side="left")
        ttk.Entry(row_alb, textvariable=self.output_name_var, width=18).pack(side="left", padx=4)
        ttk.Button(alb, text="写真帳を生成（.docx）", command=self.generate_album).pack(anchor="w", pady=2)

        # =============================================
        # RIGHT
        # =============================================
        right_wrap = ttk.Frame(main, padding=6)
        main.add(right_wrap, weight=3)

        switch_bar = ttk.Frame(right_wrap)
        switch_bar.pack(fill="x", pady=(0, 6))
        ttk.Label(switch_bar, text="プレビュー表示:").pack(side="left")
        ttk.Radiobutton(switch_bar, text="PDF", variable=self.preview_mode_var, value="pdf", command=self.update_preview_mode).pack(side="left", padx=6)
        ttk.Radiobutton(switch_bar, text="SIMA", variable=self.preview_mode_var, value="sim", command=self.update_preview_mode).pack(side="left")

        self.right_pane = ttk.Panedwindow(right_wrap, orient="vertical")
        self.right_pane.pack(fill="both", expand=True)

        # ---------------- PDF/SIMA preview area
        preview_host = ttk.Frame(self.right_pane)
        self.right_pane.add(preview_host, weight=6)

        self.pdf_preview_frame = ttk.LabelFrame(preview_host, text="PDFプレビュー", padding=6)
        self.pdf_preview_frame.pack(fill="both", expand=True)

        nav = ttk.Frame(self.pdf_preview_frame)
        nav.pack(fill="x")
        ttk.Button(nav, text="◀ 前", command=self.pdf_prev_page).pack(side="left")
        ttk.Button(nav, text="次 ▶", command=self.pdf_next_page).pack(side="left", padx=6)
        ttk.Button(nav, text="－", width=4, command=self.pdf_zoom_out).pack(side="left", padx=(12, 2))
        ttk.Button(nav, text="＋", width=4, command=self.pdf_zoom_in).pack(side="left", padx=2)
        ttk.Button(nav, text="ズームリセット", command=self.pdf_zoom_reset).pack(side="left", padx=8)
        self.pdf_page_label = ttk.Label(nav, text="page: -/-")
        self.pdf_page_label.pack(side="left", padx=10)
        self.pdf_zoom_label = ttk.Label(nav, text="zoom: 150%")
        self.pdf_zoom_label.pack(side="left", padx=10)

        pdf_canvas_frame = ttk.Frame(self.pdf_preview_frame)
        pdf_canvas_frame.pack(fill="both", expand=True)

        self.pdf_canvas = tk.Canvas(pdf_canvas_frame, bg="gray90", highlightthickness=0)
        self.pdf_canvas.pack(side="left", fill="both", expand=True)

        pdf_ys = ttk.Scrollbar(pdf_canvas_frame, orient="vertical", command=self.pdf_canvas.yview)
        pdf_ys.pack(side="right", fill="y")
        pdf_xs = ttk.Scrollbar(self.pdf_preview_frame, orient="horizontal", command=self.pdf_canvas.xview)
        pdf_xs.pack(fill="x")

        self.pdf_canvas.configure(yscrollcommand=pdf_ys.set, xscrollcommand=pdf_xs.set)
        self.pdf_canvas.bind("<ButtonPress-1>", self.on_pdf_pan_start)
        self.pdf_canvas.bind("<B1-Motion>", self.on_pdf_pan_move)
        self.pdf_canvas.bind("<ButtonRelease-1>", self.on_pdf_pan_end)
        self.pdf_canvas.bind("<Double-Button-1>", self.on_pdf_double_click_assign)
        self.pdf_canvas.bind("<Enter>", lambda e: self.pdf_canvas.focus_set())
        self.pdf_canvas.bind("<Button-1>", lambda e: self.pdf_canvas.focus_set(), add="+")
        self.pdf_canvas.bind("<MouseWheel>", self.on_pdf_wheel)
        self.pdf_canvas.bind("<Control-MouseWheel>", self.on_pdf_zoom_wheel)
        self.pdf_canvas.bind("<Button-4>", self.on_pdf_wheel_linux)
        self.pdf_canvas.bind("<Button-5>", self.on_pdf_wheel_linux)
        self.pdf_canvas.create_text(20, 20, text="（PDFを選択すると表示）", anchor="nw")

        self.sim_preview_frame = ttk.LabelFrame(preview_host, text="SIMAプレビュー（背景なし）", padding=6)

        simnav = ttk.Frame(self.sim_preview_frame)
        simnav.pack(fill="x")
        ttk.Button(simnav, text="フィット", command=self.sim_fit_view).pack(side="left")
        ttk.Button(simnav, text="ズーム100%", command=self.sim_zoom_100).pack(side="left", padx=6)
        ttk.Label(simnav, text="クリック半径(px):").pack(side="left", padx=(12, 4))
        ttk.Spinbox(simnav, from_=2, to=80, textvariable=self.sim_pick_radius_var, width=5).pack(side="left")
        ttk.Label(simnav, text="ドラッグ閾値(px):").pack(side="left", padx=(12, 4))
        ttk.Spinbox(simnav, from_=1, to=50, textvariable=self.sim_drag_threshold_var, width=5).pack(side="left")
        self.sim_info_label = ttk.Label(simnav, text="SIMA: -")
        self.sim_info_label.pack(side="left", padx=10)

        sim_canvas_frame = ttk.Frame(self.sim_preview_frame)
        sim_canvas_frame.pack(fill="both", expand=True)

        self.sim_canvas = tk.Canvas(sim_canvas_frame, bg="white", highlightthickness=0)
        self.sim_canvas.pack(side="left", fill="both", expand=True)

        sim_ys = ttk.Scrollbar(sim_canvas_frame, orient="vertical", command=self.sim_canvas.yview)
        sim_ys.pack(side="right", fill="y")
        sim_xs = ttk.Scrollbar(self.sim_preview_frame, orient="horizontal", command=self.sim_canvas.xview)
        sim_xs.pack(fill="x")

        self.sim_canvas.configure(yscrollcommand=sim_ys.set, xscrollcommand=sim_xs.set)
        self.sim_canvas.bind("<ButtonPress-1>", self.on_sim_pan_start)
        self.sim_canvas.bind("<B1-Motion>", self.on_sim_pan_move)
        self.sim_canvas.bind("<ButtonRelease-1>", self.on_sim_click_pick)
        self.sim_canvas.bind("<Double-Button-1>", self.on_sim_double_click_rename)
        self.sim_canvas.bind("<MouseWheel>", self.on_sim_zoom_wheel)
        self.sim_canvas.bind("<Control-MouseWheel>", self.on_sim_zoom_wheel)
        self.sim_canvas.bind("<Button-4>", self.on_sim_zoom_wheel_linux)
        self.sim_canvas.bind("<Button-5>", self.on_sim_zoom_wheel_linux)
        self.sim_canvas.create_text(20, 20, text="（SIMを選択すると表示）", anchor="nw")

        # ---------------- photo preview area
        photo_box = ttk.LabelFrame(self.right_pane, text="写真プレビュー（選択写真）", padding=6)
        self.right_pane.add(photo_box, weight=1)

        self.photo_preview_label = ttk.Label(photo_box, text="（写真を選択すると表示）", anchor="center")
        self.photo_preview_label.pack(fill="both", expand=True)
        self.photo_preview_label.bind("<Configure>", self.on_photo_preview_resize)

        self.update_preview_mode()

    def copy_category_choice_to_manual(self):
        self.category_manual_var.set(self.category_var.get())

    def _get_effective_category(self) -> str:
        cat = (self.category_manual_var.get() or "").strip()
        if not cat:
            cat = (self.category_var.get() or "新設").strip()
        return cat or "新設"

    def apply_category_to_selected(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("未選択", "写真一覧で区分設定する写真を選択してください。")
            return
        cat = self._get_effective_category()
        for iid in sel:
            idx = self.tree.index(iid)
            if 0 <= idx < len(self.photos):
                self.photos[idx].category = cat
        self.refresh_photo_table()
        self.logln(f"区分設定: {cat} を {len(sel)}件に設定")

    def apply_category_to_all(self):
        if not self.photos:
            return
        cat = self._get_effective_category()
        for ph in self.photos:
            ph.category = cat
        self.refresh_photo_table()
        self.logln(f"区分設定: {cat} を全写真に設定")

    def _sanitize_filename(self, s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r'[\/:*?"<>|]+', '_', s)
        s = s.rstrip('. ')
        return s

    def _get_pdf_base_image(self):
        if not self._pdf_doc:
            return None
        key = (self._pdf_page_index, round(self._pdf_zoom, 4))
        img = self._pdf_cache.get(key)
        if img is not None:
            return img
        page = self._pdf_doc[self._pdf_page_index]
        mat = fitz.Matrix(self._pdf_zoom, self._pdf_zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        self._pdf_cache[key] = img
        if len(self._pdf_cache) > 12:
            first = next(iter(self._pdf_cache))
            del self._pdf_cache[first]
        return img

    def pdf_fit_to_canvas(self):
        if not self._pdf_doc:
            return
        self.pdf_canvas.update_idletasks()
        cw = max(1, self.pdf_canvas.winfo_width())
        ch = max(1, self.pdf_canvas.winfo_height())
        page = self._pdf_doc[self._pdf_page_index]
        rect = page.rect
        zx = (cw - 20) / max(1.0, rect.width)
        zy = (ch - 20) / max(1.0, rect.height)
        self._pdf_zoom = max(self._pdf_min_zoom, min(self._pdf_max_zoom, min(zx, zy)))
        self._pdf_cache.pop((self._pdf_page_index, round(self._pdf_zoom, 4)), None)
        img = self._get_pdf_base_image()
        if img is None:
            return
        self._pdf_offx = (cw - img.width) / 2
        self._pdf_offy = (ch - img.height) / 2
        self.render_pdf_preview()

    def on_pdf_canvas_configure(self, event):
        if not self._pdf_doc:
            return
        if not self._pdf_canvas_initialized:
            self._pdf_canvas_initialized = True
            self.after(10, self.pdf_fit_to_canvas)

    # -------------------------
    # global wheel for left scroll area
    # -------------------------
    def _on_global_mousewheel(self, event):
        w = event.widget
        # Treeview・Listbox上では自前スクロールに任せる
        try:
            cls = w.winfo_class()
            if cls in ("Treeview", "Listbox"):
                return
        except Exception:
            pass
        if str(w).startswith(str(self.pdf_canvas)) or str(w).startswith(str(self.sim_canvas)):
            return
        try:
            self.left_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        except Exception:
            pass

    def _on_global_mousewheel_linux(self, event):
        w = event.widget
        try:
            cls = w.winfo_class()
            if cls in ("Treeview", "Listbox"):
                return
        except Exception:
            pass
        if str(w).startswith(str(self.pdf_canvas)) or str(w).startswith(str(self.sim_canvas)):
            return
        try:
            if event.num == 4:
                self.left_canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                self.left_canvas.yview_scroll(1, "units")
        except Exception:
            pass

    def update_preview_mode(self):
        mode = self.preview_mode_var.get()
        if mode == "pdf":
            self.sim_preview_frame.pack_forget()
            self.pdf_preview_frame.pack(fill="both", expand=True)
        else:
            self.pdf_preview_frame.pack_forget()
            self.sim_preview_frame.pack(fill="both", expand=True)

    # -------------------------
    # photos
    # -------------------------
    def add_photos_dialog(self):
        paths = filedialog.askopenfilenames(
            title="写真を選択",
            filetypes=[("Images", "*.jpg;*.jpeg;*.png;*.bmp;*.tif;*.tiff;*.webp"), ("All files", "*.*")]
        )
        if paths:
            self.add_photos(list(paths))

    def on_drop_photos(self, event):
        paths = norm_paths_from_dnd(event.data)
        imgs = [p for p in paths if os.path.isfile(p) and is_image_file(p)]
        if not imgs:
            self.logln("画像ファイルが見つかりませんでした。")
            return
        self.add_photos(imgs)

    def add_photos(self, paths: List[str]):
        added = 0
        existing = {os.path.abspath(p.src_path) for p in self.photos}
        for p in paths:
            ap = os.path.abspath(p)
            if os.path.isfile(ap) and is_image_file(ap) and ap not in existing:
                self.photos.append(PhotoItem(src_path=ap, category=self._get_effective_category()))
                existing.add(ap)
                added += 1
        self.refresh_photo_table()
        self.logln(f"写真追加: {added}件")

    def clear_photos(self):
        self.photos = []
        self.refresh_photo_table()
        self._photo_preview_tk = None
        self._current_photo_preview_path = ""
        self.photo_preview_label.configure(image="", text="（写真を選択すると表示）")
        self.logln("写真一覧をクリアしました")

    def _on_tree_click(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region != "cell":
            return
        col = self.tree.identify_column(event.x)
        row = self.tree.identify_row(event.y)
        if not row:
            return
        # 区分列は#4（説明列は#5）
        if col != "#4":
            return
        self.after(10, lambda: self._show_category_combobox(row, col))

    def _on_tree_dblclick(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region != "cell":
            return
        col = self.tree.identify_column(event.x)
        row = self.tree.identify_row(event.y)
        if not row:
            return
        if col == "#5":  # 説明列
            self.after(10, lambda: self._show_note_entry(row, col))

    def _show_note_entry(self, row_iid, col):
        bbox = self.tree.bbox(row_iid, col)
        if not bbox:
            return
        x, y, width, height = bbox

        idx = self.tree.index(row_iid)
        if idx < 0 or idx >= len(self.photos):
            return
        current_val = self.photos[idx].note or ""

        var = tk.StringVar(value=current_val)
        entry = ttk.Entry(self.tree, textvariable=var)
        entry.place(x=x, y=y, width=max(width, 160), height=height + 2)
        entry.focus_set()
        entry.select_range(0, "end")

        committed = [False]

        def commit(e=None):
            if committed[0]:
                return
            committed[0] = True
            val = var.get()
            if 0 <= idx < len(self.photos):
                self.photos[idx].note = val
                self.refresh_photo_table()
                self.logln(f"説明変更: No.{idx+1} → {val}")
            try:
                entry.destroy()
            except Exception:
                pass

        def on_escape(e=None):
            committed[0] = True
            try:
                entry.destroy()
            except Exception:
                pass

        entry.bind("<Return>", commit)
        entry.bind("<Tab>", commit)
        entry.bind("<FocusOut>", commit)
        entry.bind("<Escape>", on_escape)

    def _show_category_combobox(self, row_iid, col):
        bbox = self.tree.bbox(row_iid, col)
        if not bbox:
            return
        x, y, width, height = bbox

        idx = self.tree.index(row_iid)
        if idx < 0 or idx >= len(self.photos):
            return
        current_val = self.photos[idx].category or ""

        var = tk.StringVar(value=current_val)
        cb = ttk.Combobox(self.tree, textvariable=var,
                          values=["新設", "復元", "既設"],
                          width=8)
        cb.place(x=x, y=y, width=max(width, 90), height=height + 2)
        cb.focus_set()
        cb.select_range(0, "end")

        committed = [False]

        def commit(e=None):
            if committed[0]:
                return
            committed[0] = True
            val = var.get().strip()
            if not val:
                val = "新設"
            if 0 <= idx < len(self.photos):
                self.photos[idx].category = val
                self.refresh_photo_table()
                self.logln(f"区分変更: No.{idx+1} → {val}")
            try:
                cb.destroy()
            except Exception:
                pass

        def on_escape(e=None):
            committed[0] = True
            try:
                cb.destroy()
            except Exception:
                pass

        cb.bind("<<ComboboxSelected>>", commit)
        cb.bind("<Return>", commit)
        cb.bind("<Tab>", commit)
        cb.bind("<FocusOut>", commit)
        cb.bind("<Escape>", on_escape)

    def refresh_photo_table(self):
        self.tree.delete(*self.tree.get_children())
        for i, ph in enumerate(self.photos, start=1):
            src = os.path.basename(ph.src_path)
            pt = ph.assigned_point
            rz = os.path.basename(ph.resized_path) if ph.resized_path else ""
            rn = os.path.basename(ph.renamed_path) if ph.renamed_path else ""
            cat = ph.category or ""
            note = ph.note or ""
            self.tree.insert("", "end", values=(i, src, pt, cat, note, rz, rn))

    def sort_photos_by_exif(self):
        if not self.photos:
            return
        self.photos.sort(key=lambda ph: get_exif_datetime_key(ph.src_path))
        self.refresh_photo_table()
        self.logln("EXIFで並び替えました")

    def on_photo_select_preview(self, event=None):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        if idx < 0 or idx >= len(self.photos):
            return
        self.show_photo_preview(self.photos[idx].src_path)

    def on_photo_preview_resize(self, event=None):
        if self._current_photo_preview_path and os.path.isfile(self._current_photo_preview_path):
            self.show_photo_preview(self._current_photo_preview_path)

    def show_photo_preview(self, path: str):
        try:
            self._current_photo_preview_path = path
            im = Image.open(path)
            im = ImageOps.exif_transpose(im)

            self.photo_preview_label.update_idletasks()
            w = max(500, self.photo_preview_label.winfo_width() or 700)
            h = max(300, self.photo_preview_label.winfo_height() or 420)

            im.thumbnail((w, h), Image.LANCZOS)

            tkimg = ImageTk.PhotoImage(im)
            self._photo_preview_tk = tkimg
            self.photo_preview_label.configure(image=tkimg, text="")
        except Exception as e:
            self.photo_preview_label.configure(image="", text=f"プレビュー失敗: {e}")
            self._photo_preview_tk = None

    # -------------------------
    # points
    # -------------------------
    def set_point_list(self, pts: List[str]):
        pts2 = [p for p in pts if (p or "").strip()]
        self.point_names = _dedupe_keep_order([normalize_point(p) for p in pts2])
        self.points_list.delete(0, "end")
        for p in self.point_names:
            self.points_list.insert("end", p)

    def _sync_points_from_listbox(self):
        self.point_names = [self.points_list.get(i) for i in range(self.points_list.size())]

    def on_point_select(self, event=None):
        sel = self.points_list.curselection()
        if not sel:
            return
        idx = sel[0]
        val = self.points_list.get(idx)
        self.point_edit_var.set(val)
        self._pdf_highlight_word = val
        self._sim_selected = val
        self.render_pdf_preview()
        self.render_sim_view()

    def point_add(self):
        v = normalize_point(self.point_edit_var.get())
        if not v:
            return
        existing = set(self.points_list.get(i) for i in range(self.points_list.size()))
        if v in existing:
            messagebox.showinfo("重複", "同じ点名が既にあります。")
            return
        sel = self.points_list.curselection()
        if sel:
            ins = sel[0] + 1
            self.points_list.insert(ins, v)
        else:
            self.points_list.insert("end", v)
        self._sync_points_from_listbox()
        self.logln(f"点名追加: {v}")

    def point_update(self):
        sel = self.points_list.curselection()
        if not sel:
            messagebox.showinfo("未選択", "更新したい点名を一覧から選択してください。")
            return
        v = normalize_point(self.point_edit_var.get())
        if not v:
            return
        idx = sel[0]
        self.points_list.delete(idx)
        self.points_list.insert(idx, v)
        self.points_list.selection_set(idx)
        self._sync_points_from_listbox()
        self._pdf_highlight_word = v
        self._sim_selected = v
        self.render_pdf_preview()
        self.render_sim_view()
        self.logln(f"点名更新: {v}")

    def point_delete(self):
        sel = list(self.points_list.curselection())
        if not sel:
            return
        for idx in reversed(sel):
            self.points_list.delete(idx)
        self._sync_points_from_listbox()
        self.point_edit_var.set("")
        self._pdf_highlight_word = ""
        self._sim_selected = ""
        self.render_pdf_preview()
        self.render_sim_view()
        self.logln(f"点名削除: {len(sel)}件")

    def point_move_up(self):
        sel = self.points_list.curselection()
        if not sel:
            return
        idx = sel[0]
        if idx <= 0:
            return
        val = self.points_list.get(idx)
        self.points_list.delete(idx)
        self.points_list.insert(idx - 1, val)
        self.points_list.selection_set(idx - 1)
        self._sync_points_from_listbox()

    def point_move_down(self):
        sel = self.points_list.curselection()
        if not sel:
            return
        idx = sel[0]
        if idx >= self.points_list.size() - 1:
            return
        val = self.points_list.get(idx)
        self.points_list.delete(idx)
        self.points_list.insert(idx + 1, val)
        self.points_list.selection_set(idx + 1)
        self._sync_points_from_listbox()

    # -------------------------
    # assign
    # -------------------------
    def assign_points_in_order(self):
        photo_sel = self.tree.selection()
        if not photo_sel:
            messagebox.showinfo("未選択", "写真一覧で割当対象の行を選択してください。")
            return
        pt_sel = list(self.points_list.curselection())
        if not pt_sel:
            messagebox.showinfo("未選択", "点名一覧で割当対象の点名を選択してください。")
            return

        sel_indices = sorted([self.tree.index(iid) for iid in photo_sel])
        sel_points = [self.point_names[i] for i in pt_sel]

        n = min(len(sel_indices), len(sel_points))
        for k in range(n):
            self.photos[sel_indices[k]].assigned_point = sel_points[k]

        self.refresh_photo_table()
        self.logln(f"手動割当（順番）: {n}件")

    def on_point_double_click_assign(self, event=None):
        sel_pt = self.points_list.curselection()
        if not sel_pt:
            return
        point = (self.points_list.get(sel_pt[0]) or "").strip()
        if not point:
            return

        photo_sel = self.tree.selection()
        if not photo_sel:
            messagebox.showinfo("写真未選択", "左の写真一覧で割当対象の写真を選択してください。")
            return

        sel_indices = sorted([self.tree.index(iid) for iid in photo_sel])
        for idx in sel_indices:
            self.photos[idx].assigned_point = point

        self.refresh_photo_table()
        self.logln(f"ダブルクリック割当: {point} を {len(sel_indices)}件に設定")

    def clear_selected_assignments(self):
        photo_sel = self.tree.selection()
        if not photo_sel:
            return
        sel_indices = [self.tree.index(iid) for iid in photo_sel]
        for idx in sel_indices:
            self.photos[idx].assigned_point = ""
        self.refresh_photo_table()
        self.logln(f"点名クリア: {len(sel_indices)}件")

    def auto_assign_points_1to1(self):
        if not self.photos:
            messagebox.showinfo("写真なし", "写真を追加してください。")
            return
        if not self.point_names:
            messagebox.showinfo("点名なし", "先に点名抽出（PDF or SIM）をしてください。")
            return

        n = min(len(self.photos), len(self.point_names))
        for i in range(n):
            self.photos[i].assigned_point = self.point_names[i]
        for i in range(n, len(self.photos)):
            self.photos[i].assigned_point = ""

        self.refresh_photo_table()
        self.logln(f"自動割当: {n}件割当、残り{len(self.photos)-n}件は未割当")

    def sort_photos_by_point_order(self):
        if not self.photos:
            return
        if not self.point_names:
            messagebox.showinfo("点名なし", "先に点名抽出（PDF or SIM）をしてください。")
            return

        order = {p: i for i, p in enumerate(self.point_names)}

        def keyfunc(ph: PhotoItem):
            p = (ph.assigned_point or "").strip()
            if not p:
                return (2, 10**9, "")
            idx = order.get(p, 10**8)
            return (0, idx, p)

        self.photos.sort(key=keyfunc)
        self.refresh_photo_table()
        self.logln("① 点名順で並び替えました")

    # -------------------------
    # PDF
    # -------------------------
    def pick_pdf_dialog(self):
        p = filedialog.askopenfilename(title="PDFを選択", filetypes=[("PDF", "*.pdf"), ("All files", "*.*")])
        if p:
            self.load_pdf(p)

    def on_drop_pdf(self, event):
        paths = norm_paths_from_dnd(event.data)
        pdfs = [p for p in paths if os.path.isfile(p) and os.path.splitext(p.lower())[1] == ".pdf"]
        if not pdfs:
            self.logln("PDFが見つかりませんでした。")
            return
        self.load_pdf(pdfs[0])

    def load_pdf(self, path: str):
        self.current_pdf = path
        self.pdf_label.config(text=f"PDF: {os.path.basename(path)}")
        self.logln(f"PDF選択: {path}")
        try:
            if self._pdf_doc:
                self._pdf_doc.close()
            self._pdf_doc = fitz.open(path)
            self._pdf_page_index = 0
            self._pdf_cache = {}
            self._pdf_canvas_initialized = False
            self._pdf_zoom = 1.0
            self._pdf_offx = 20.0
            self._pdf_offy = 20.0
            self.pdf_fit_to_canvas()
        except Exception as e:
            self._pdf_doc = None
            self.pdf_canvas.delete("all")
            self.pdf_canvas.create_text(20, 20, text=f"PDF読込失敗: {e}", anchor="nw")
            self.logln(f"PDF読込失敗: {e}")

    def do_extract_points_pdf(self):
        if not self.current_pdf or not os.path.isfile(self.current_pdf):
            messagebox.showwarning("PDF未選択", "PDFを選択してください。")
            return
        try:
            names = extract_point_names_from_pdf(self.current_pdf)
            self.set_point_list(names)
            self.logln(f"PDF点名抽出: {len(names)}件")
        except Exception as e:
            messagebox.showerror("PDF点名抽出エラー", str(e))
            self.logln(f"PDF点名抽出エラー: {e}")

    def pdf_prev_page(self):
        if not self._pdf_doc:
            return
        self._pdf_page_index = max(0, self._pdf_page_index - 1)
        self.pdf_fit_to_canvas()

    def pdf_next_page(self):
        if not self._pdf_doc:
            return
        self._pdf_page_index = min(len(self._pdf_doc) - 1, self._pdf_page_index + 1)
        self.pdf_fit_to_canvas()

    def pdf_zoom_in(self):
        if not self._pdf_doc:
            return
        self.pdf_zoom_at(self.pdf_canvas.winfo_width() / 2, self.pdf_canvas.winfo_height() / 2, 1.15)

    def pdf_zoom_out(self):
        if not self._pdf_doc:
            return
        self.pdf_zoom_at(self.pdf_canvas.winfo_width() / 2, self.pdf_canvas.winfo_height() / 2, 1 / 1.15)

    def pdf_zoom_reset(self):
        self.pdf_fit_to_canvas()

    def pdf_zoom_at(self, cx: float, cy: float, factor: float):
        if not self._pdf_doc:
            return
        old_zoom = self._pdf_zoom
        new_zoom = max(self._pdf_min_zoom, min(self._pdf_max_zoom, old_zoom * factor))
        if abs(new_zoom - old_zoom) < 1e-9:
            return
        self._pdf_offx = cx - (cx - self._pdf_offx) * (new_zoom / old_zoom)
        self._pdf_offy = cy - (cy - self._pdf_offy) * (new_zoom / old_zoom)
        self._pdf_zoom = new_zoom
        self.render_pdf_preview()

    def on_pdf_pan_start(self, event):
        if not self._pdf_doc:
            return
        self.pdf_canvas.focus_set()
        self._pdf_pan_start  = (event.x, event.y)
        self._pdf_pan_origin = (event.x, event.y)
        self._pdf_dragging   = True
        self._pdf_pan_moved  = False

    def on_pdf_pan_move(self, event):
        if not self._pdf_doc or not self._pdf_pan_start:
            return
        x0, y0 = self._pdf_pan_start
        ox, oy = self._pdf_pan_origin
        if abs(event.x - ox) > 4 or abs(event.y - oy) > 4:
            self._pdf_pan_moved = True
        self._pdf_offx += event.x - x0
        self._pdf_offy += event.y - y0
        self._pdf_pan_start = (event.x, event.y)
        self.render_pdf_preview()

    def on_pdf_pan_end(self, event):
        moved = self._pdf_pan_moved
        self._pdf_pan_start  = None
        self._pdf_pan_origin = None
        self._pdf_dragging   = False
        self._pdf_pan_moved  = False
        if not moved:
            self.on_pdf_click_pick(event)

    def _pdf_nearest_point_name(self, canvas_x: float, canvas_y: float) -> str:
        """クリック位置（キャンバス座標）に最も近い point_names の点名を返す。見つからなければ空文字。"""
        if not self._pdf_doc or not self.point_names:
            return ""
        # キャンバス座標 → PDFページ座標
        pdf_x = (canvas_x - self._pdf_offx) / max(self._pdf_zoom, 0.01)
        pdf_y = (canvas_y - self._pdf_offy) / max(self._pdf_zoom, 0.01)

        try:
            page = self._pdf_doc[self._pdf_page_index]
            words = page.get_text("words")  # (x0,y0,x1,y1,word,block,line,word_no)
        except Exception:
            return ""

        # point_names に含まれるワードのみ候補に絞る
        # 点名は複数ページにわたる場合もあるので全ページ検索はせず現ページのみ
        THRESHOLD = 40 / max(self._pdf_zoom, 0.01)  # canvas 40px 分のPDF座標距離

        best_name = ""
        best_dist = float("inf")
        for (x0, y0, x1, y1, word, *_) in words:
            w = normalize_point(word)
            if w not in self.point_names:
                continue
            cx = (x0 + x1) / 2
            cy = (y0 + y1) / 2
            dist = ((cx - pdf_x) ** 2 + (cy - pdf_y) ** 2) ** 0.5
            if dist < best_dist and dist < THRESHOLD:
                best_dist = dist
                best_name = w
        return best_name

    def on_pdf_click_pick(self, event):
        """PDFシングルクリック → 最近傍点名をハイライト"""
        name = self._pdf_nearest_point_name(event.x, event.y)
        if not name:
            return
        self._pdf_highlight_word = name
        self._sim_selected = name
        # 点名一覧で選択
        if name in self.point_names:
            idx = self.point_names.index(name)
            self.points_list.selection_clear(0, "end")
            self.points_list.selection_set(idx)
            self.points_list.see(idx)
            self.point_edit_var.set(name)
        self.render_pdf_preview()
        self.render_sim_view()
        self.logln(f"PDFクリック選択: {name}")

    def on_pdf_double_click_assign(self, event):
        """PDFダブルクリック → 最近傍点名を選択写真に割当（SIMAダブルクリックと同じ）"""
        name = self._pdf_nearest_point_name(event.x, event.y)
        if not name:
            return
        photo_sel = self.tree.selection()
        if not photo_sel:
            messagebox.showinfo("写真未選択", "左の写真一覧で割当対象の写真を選択してください。")
            return
        sel_indices = sorted([self.tree.index(iid) for iid in photo_sel])
        for idx in sel_indices:
            self.photos[idx].assigned_point = name
        self.refresh_photo_table()
        # ハイライトも更新
        self._pdf_highlight_word = name
        self._sim_selected = name
        self.render_pdf_preview()
        self.render_sim_view()
        self.logln(f"PDFダブルクリック割当: {name} を {len(sel_indices)}件に設定")

    def on_pdf_wheel(self, event):
        if not self._pdf_doc:
            return "break"
        ctrl = self._event_has_ctrl(event)
        shift = self._event_has_shift(event)
        if ctrl:
            factor = 1.10 if event.delta > 0 else (1 / 1.10)
            self.pdf_zoom_at(event.x, event.y, factor)
            return "break"
        step = 40 if event.delta > 0 else -40
        if shift:
            self._pdf_offx += step
        else:
            self._pdf_offy += step
        self.render_pdf_preview()
        return "break"

    def on_pdf_wheel_linux(self, event):
        if not self._pdf_doc:
            return "break"
        ctrl = self._event_has_ctrl(event)
        shift = self._event_has_shift(event)
        up = (event.num == 4)
        if ctrl:
            self.pdf_zoom_at(event.x, event.y, 1.10 if up else (1 / 1.10))
            return "break"
        step = 40 if up else -40
        if shift:
            self._pdf_offx += step
        else:
            self._pdf_offy += step
        self.render_pdf_preview()
        return "break"

    def on_pdf_zoom_wheel(self, event):
        return self.on_pdf_wheel(event)

    def render_pdf_preview(self):
        self.pdf_canvas.delete("all")

        if not self._pdf_doc:
            self.pdf_canvas.create_text(20, 20, text="（PDFを選択すると表示）", anchor="nw")
            self.pdf_page_label.config(text="page: -/-")
            try:
                self.pdf_zoom_label.config(text="zoom: -")
            except Exception:
                pass
            return

        self.pdf_page_label.config(text=f"page: {self._pdf_page_index + 1}/{len(self._pdf_doc)}")
        try:
            self.pdf_zoom_label.config(text=f"zoom: {int(self._pdf_zoom * 100)}%")
        except Exception:
            pass

        base_img = self._get_pdf_base_image()
        if base_img is None:
            return
        img = base_img.copy()

        word = (self._pdf_highlight_word or "").strip()
        if word:
            try:
                page = self._pdf_doc[self._pdf_page_index]
                rects = page.search_for(word)
                if rects:
                    draw = ImageDraw.Draw(img)
                    for r in rects:
                        x0 = r.x0 * self._pdf_zoom
                        y0 = r.y0 * self._pdf_zoom
                        x1 = r.x1 * self._pdf_zoom
                        y1 = r.y1 * self._pdf_zoom
                        draw.rectangle([x0, y0, x1, y1], outline=(255, 0, 0), width=3)
            except Exception:
                pass

        tkimg = ImageTk.PhotoImage(img)
        self._pdf_preview_tk = tkimg
        self.pdf_canvas.create_image(self._pdf_offx, self._pdf_offy, image=tkimg, anchor="nw")
        self.pdf_canvas.create_text(10, 10, text="Ctrl+ホイール: マウス中心ズーム / 左ドラッグ: パン", anchor="nw", fill="gray25")

    # -------------------------
    # SIMA
    # -------------------------
    def pick_sim_dialog(self):
        p = filedialog.askopenfilename(title="SIMA(.sim/.sima)を選択", filetypes=[("SIMA", "*.sim;*.sima"), ("All files", "*.*")])
        if p:
            self.load_sim(p)

    def on_drop_sim(self, event):
        paths = norm_paths_from_dnd(event.data)
        sims = [p for p in paths if os.path.isfile(p) and os.path.splitext(p.lower())[1] in (".sim", ".sima")]
        if not sims:
            self.logln("SIMファイルが見つかりませんでした。")
            return
        self.load_sim(sims[0])

    def load_sim(self, path: str):
        self.current_sim = path
        self.sim_label.config(text=f"SIMA: {os.path.basename(path)}")
        self.logln(f"SIMA選択: {path}")
        try:
            pts, polys = parse_sima_file(path)
            self.sima_points = pts
            self.sima_polylines = polys
            self._sim_selected = ""
            self._compute_sim_bbox()
            self.sim_fit_view()
        except Exception as e:
            self.sima_points = {}
            self.sima_polylines = []
            self.sim_canvas.delete("all")
            self.sim_canvas.create_text(20, 20, text=f"SIMA読込失敗: {e}", anchor="nw")
            self.logln(f"SIMA読込失敗: {e}")

    def do_extract_points_sim(self):
        if not self.current_sim or not os.path.isfile(self.current_sim):
            messagebox.showwarning("SIM未選択", "SIMA(.sim/.sima)を選択してください。")
            return
        try:
            pts, polys = parse_sima_file(self.current_sim)
            self.sima_points = pts
            self.sima_polylines = polys
            self.set_point_list(list(pts.keys()))
            self._compute_sim_bbox()
            self.sim_fit_view()
            self.logln(f"SIM点名抽出: {len(pts)}件")
        except Exception as e:
            messagebox.showerror("SIM点名抽出エラー", str(e))
            self.logln(f"SIM点名抽出エラー: {e}")

    def _compute_sim_bbox(self):
        if not self.sima_points:
            self._sim_bbox = None
            return

        # X=北, Y=東西
        xs = [v[0] for v in self.sima_points.values()]
        ys = [v[1] for v in self.sima_points.values()]
        self._sim_bbox = (min(xs), min(ys), max(xs), max(ys))

    def sim_zoom_100(self):
        self._sim_zoom = 1.0
        self._sim_offx = 40
        self._sim_offy = 40
        self.render_sim_view()

    def sim_fit_view(self):
        self.sim_canvas.update_idletasks()
        cw = max(1, self.sim_canvas.winfo_width())
        ch = max(1, self.sim_canvas.winfo_height())

        if not self._sim_bbox:
            self.sim_canvas.delete("all")
            self.sim_canvas.create_text(20, 20, text="（SIMを選択すると表示）", anchor="nw")
            self.sim_info_label.config(text="SIMA: -")
            return

        min_n, min_e, max_n, max_e = self._sim_bbox

        # 横 = Y(東西), 縦 = X(北)
        dw = max(1e-9, max_e - min_e)
        dh = max(1e-9, max_n - min_n)

        margin = 40
        sx = (cw - 2 * margin) / dw
        sy = (ch - 2 * margin) / dh
        self._sim_zoom = max(self._sim_min_zoom, min(self._sim_max_zoom, min(sx, sy)))

        self._sim_offx = margin - min_e * self._sim_zoom
        self._sim_offy = margin + max_n * self._sim_zoom

        self.render_sim_view()

    def world_to_screen(self, north_x: float, east_y: float) -> Tuple[float, float]:
        # 横 = 東西(Y), 縦 = 北(X)
        sx = east_y * self._sim_zoom + self._sim_offx
        sy = -north_x * self._sim_zoom + self._sim_offy
        return sx, sy

    def screen_to_world(self, sx: float, sy: float) -> Tuple[float, float]:
        east_y = (sx - self._sim_offx) / self._sim_zoom
        north_x = -(sy - self._sim_offy) / self._sim_zoom
        return north_x, east_y

    def on_sim_pan_start(self, event):
        self._sim_pan_start = (event.x, event.y)
        self._sim_pan_origin = (event.x, event.y)
        self._sim_pan_moved = False

    def on_sim_pan_move(self, event):
        if not self._sim_pan_start:
            return
        x0, y0 = self._sim_pan_start
        dx = event.x - x0
        dy = event.y - y0

        ox, oy = self._sim_pan_origin if self._sim_pan_origin else (event.x, event.y)
        moved = math.hypot(event.x - ox, event.y - oy)
        if moved >= int(self.sim_drag_threshold_var.get() or 0):
            self._sim_pan_moved = True

        self._sim_offx += dx
        self._sim_offy += dy
        self._sim_pan_start = (event.x, event.y)
        self.render_sim_view()

    def on_sim_click_pick(self, event):
        if not self.sima_points:
            return
        if self._sim_pan_moved:
            self._sim_pan_start = None
            self._sim_pan_origin = None
            self._sim_pan_moved = False
            return

        px, py = event.x, event.y
        rad = int(self.sim_pick_radius_var.get() or 0)
        if rad <= 0:
            return

        best_name = None
        best_d2 = rad * rad

        for name, (north_x, east_y, _) in self.sima_points.items():
            sx, sy = self.world_to_screen(north_x, east_y)
            dx = sx - px
            dy = sy - py
            d2 = dx * dx + dy * dy
            if d2 <= best_d2:
                best_d2 = d2
                best_name = name

        self._sim_pan_start = None
        self._sim_pan_origin = None
        self._sim_pan_moved = False

        if not best_name:
            return

        bn = normalize_point(best_name)
        if bn not in self.point_names:
            self.point_names.append(bn)
            self.points_list.insert("end", bn)

        idx = self.point_names.index(bn)
        self.points_list.selection_clear(0, "end")
        self.points_list.selection_set(idx)
        self.points_list.see(idx)
        self.point_edit_var.set(bn)

        self._sim_selected = bn
        self._pdf_highlight_word = bn
        self.render_sim_view()
        self.render_pdf_preview()
        self.logln(f"SIMAクリック選択: {bn}")

    def on_sim_double_click_rename(self, event):
        """SIMAキャンバス上のダブルクリック → 最近傍点名を選択写真に割当（点名一覧ダブルクリックと同じ）"""
        if not self.sima_points:
            return

        px = self.sim_canvas.canvasx(event.x)
        py = self.sim_canvas.canvasy(event.y)
        rad = max(30, int(self.sim_pick_radius_var.get() or 0) * 3)

        best_name = None
        best_d2 = rad * rad
        for name, (north_x, east_y, _) in self.sima_points.items():
            sx, sy = self.world_to_screen(north_x, east_y)
            d2 = (sx - px) ** 2 + (sy - py) ** 2
            if d2 <= best_d2:
                best_d2 = d2
                best_name = name

        if not best_name:
            return

        point = normalize_point(best_name)

        photo_sel = self.tree.selection()
        if not photo_sel:
            messagebox.showinfo("写真未選択", "左の写真一覧で割当対象の写真を選択してください。")
            return

        sel_indices = sorted([self.tree.index(iid) for iid in photo_sel])
        for idx in sel_indices:
            self.photos[idx].assigned_point = point

        self.refresh_photo_table()
        self.logln(f"SIMAダブルクリック割当: {point} を {len(sel_indices)}件に設定")

    def on_sim_zoom_wheel(self, event):
        if not self.sima_points:
            return
        mx, my = event.x, event.y
        wx0, wy0 = self.screen_to_world(mx, my)

        if event.delta > 0:
            self._sim_zoom *= 1.10
        else:
            self._sim_zoom /= 1.10
        self._sim_zoom = max(self._sim_min_zoom, min(self._sim_max_zoom, self._sim_zoom))

        sx1, sy1 = self.world_to_screen(wx0, wy0)
        self._sim_offx += (mx - sx1)
        self._sim_offy += (my - sy1)
        self.render_sim_view()

    def on_sim_zoom_wheel_linux(self, event):
        if not self.sima_points:
            return
        mx, my = event.x, event.y
        wx0, wy0 = self.screen_to_world(mx, my)

        if event.num == 4:
            self._sim_zoom *= 1.10
        elif event.num == 5:
            self._sim_zoom /= 1.10
        self._sim_zoom = max(self._sim_min_zoom, min(self._sim_max_zoom, self._sim_zoom))

        sx1, sy1 = self.world_to_screen(wx0, wy0)
        self._sim_offx += (mx - sx1)
        self._sim_offy += (my - sy1)
        self.render_sim_view()

    def render_sim_view(self):
        self.sim_canvas.delete("all")

        if not self.sima_points:
            self.sim_canvas.create_text(20, 20, text="（SIMを選択すると表示）", anchor="nw")
            self.sim_info_label.config(text="SIMA: -")
            return

        min_n, min_e, max_n, max_e = self._sim_bbox if self._sim_bbox else (0, 0, 0, 0)

        p1 = self.world_to_screen(min_n, min_e)
        p2 = self.world_to_screen(max_n, max_e)
        left_s = min(p1[0], p2[0]) - 200
        right_s = max(p1[0], p2[0]) + 200
        top_s = min(p1[1], p2[1]) - 200
        bottom_s = max(p1[1], p2[1]) + 200
        self.sim_canvas.config(scrollregion=(left_s, top_s, right_s, bottom_s))

        for poly in self.sima_polylines:
            pts = []
            for name in poly:
                if name in self.sima_points:
                    north_x, east_y, _ = self.sima_points[name]
                    sx, sy = self.world_to_screen(north_x, east_y)
                    pts.extend([sx, sy])
            if len(pts) >= 4:
                self.sim_canvas.create_line(*pts, fill="black", width=2)
                if len(pts) >= 6:
                    self.sim_canvas.create_line(pts[-2], pts[-1], pts[0], pts[1], fill="black", width=2)

        r = 3
        for name, (north_x, east_y, _) in self.sima_points.items():
            sx, sy = self.world_to_screen(north_x, east_y)
            n = normalize_point(name)
            fill = "red" if n == self._sim_selected else "black"
            self.sim_canvas.create_oval(sx - r, sy - r, sx + r, sy + r, fill=fill, outline=fill)
            self.sim_canvas.create_text(sx + 6, sy - 6, text=n, anchor="sw", fill=fill, font=("Meiryo", 10))

        self.sim_info_label.config(text=f"SIMA: 点{len(self.sima_points)} / 画地{len(self.sima_polylines)}  zoom:{self._sim_zoom:.3f}x")

    # -------------------------
    # rename / resize
    # -------------------------
    def do_resize(self):
        if not self.photos:
            messagebox.showinfo("写真なし", "写真を追加してください。")
            return
        done = 0
        for ph in self.photos:
            src = ph.src_path
            dst_dir = os.path.join(os.path.dirname(src), "resize")
            dst = os.path.join(dst_dir, os.path.basename(src))
            try:
                resize_keep_aspect(src, dst, TARGET_W, TARGET_H)
                ph.resized_path = dst
                done += 1
            except Exception as e:
                self.logln(f"リサイズ失敗: {src} / {e}")
        self.refresh_photo_table()
        self.logln(f"リサイズ完了: {done}件")

    def do_rename_global_prefix(self):
        if not self.photos:
            messagebox.showinfo("写真なし", "写真を追加してください。")
            return

        missing = [ph for ph in self.photos if not ph.resized_path or not os.path.isfile(ph.resized_path)]
        if missing:
            messagebox.showwarning("未リサイズあり", "先にリサイズしてください。")
            return

        mode = self.rename_mode_var.get()
        done = 0

        if mode == "point":
            no_point = [i for i, ph in enumerate(self.photos, start=1) if not (ph.assigned_point or "").strip()]
            if no_point:
                messagebox.showwarning("点名未割当", f"点名が未割当の写真があります（No: {no_point[:30]}...）")
                return

        start_no = int(self.serial_start_var.get() or 1)
        digits = int(self.serial_digits_var.get() or 1)
        prefix = self.serial_prefix_var.get() or ""
        use_fullwidth = bool(self.serial_fullwidth_var.get())

        for i, ph in enumerate(self.photos, start=1):
            src = ph.resized_path
            folder = os.path.dirname(src)
            ext = os.path.splitext(src)[1].lower()

            if mode == "point":
                point = (ph.assigned_point or "").strip()
                new_name = f"{i}_{point}{ext}"
            else:
                serial_no = start_no + (i - 1)
                num = str(serial_no).zfill(digits)
                if use_fullwidth:
                    num = to_fullwidth_digits(num)
                new_name = f"{prefix}{num}{ext}"

            dst = os.path.join(folder, new_name)
            if os.path.exists(dst):
                base, ext2 = os.path.splitext(new_name)
                dst = os.path.join(folder, f"{base}_{datetime.now().strftime('%H%M%S')}{ext2}")

            try:
                os.rename(src, dst)
                ph.renamed_path = dst
                ph.resized_path = dst
                done += 1
            except Exception as e:
                self.logln(f"リネーム失敗: {src} / {e}")

        self.refresh_photo_table()

        if mode == "point":
            self.logln(f"リネーム完了: {done}件（点名付き）")
        else:
            self.logln(
                f"リネーム完了: {done}件（連番のみ / 接頭語='{prefix}' / 開始={start_no} / 桁数={digits} / 全角={use_fullwidth}）"
            )

    # -------------------------
    # album
    # -------------------------
    def generate_album(self):
        if not self.photos:
            messagebox.showinfo("写真なし", "写真を追加してください。")
            return

        out_dir = os.path.dirname(self.photos[0].src_path)
        safe_mkdir(out_dir)

        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = self._sanitize_filename(self.output_name_var.get())
        if not base_name:
            base_name = f"写真帳_{stamp}"
        if not base_name.lower().endswith(".docx"):
            base_name += ".docx"
        out_docx = os.path.join(out_dir, base_name)

        rows = []
        for ph in self.photos:
            img_path = ph.renamed_path or ph.resized_path or ph.src_path
            if not os.path.isfile(img_path):
                continue
            name_no_ext = os.path.splitext(os.path.basename(img_path))[0]
            shoot = get_shoot_date_jp(ph.src_path)
            rows.append((
                img_path,
                name_no_ext,
                shoot,
                ph.category or self._get_effective_category(),
                ph.note or "",
            ))

        if not rows:
            messagebox.showwarning("画像なし", "写真帳に入れられる画像がありません。")
            return

        try:
            build_album_direct(out_docx, rows, photos_per_page=4, img_width_mm=80)
            self.logln(f"写真帳生成: {out_docx}")
            messagebox.showinfo("完了", f"写真帳を生成しました。\n{out_docx}")
        except Exception as e:
            messagebox.showerror("写真帳生成エラー", str(e))
            self.logln(f"写真帳生成エラー: {e}")


def main():
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()